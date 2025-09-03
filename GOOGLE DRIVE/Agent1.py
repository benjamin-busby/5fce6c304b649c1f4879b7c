"""
Agent1 — CV Gatekeeper (GUI-friendly)
- Scans a folder for .pdf, .docx, .txt files
- Robust text extraction (PDF fallbacks; DOCX incl. tables/headers/text boxes; legacy .doc auto-convert)
- Order of checks: LENGTH → LANGUAGE → CV HEURISTIC
- Streams print() to GUI via run_from_gui(callback)
- Exposes bad_files via run_and_get_bad_files(callback=None)
"""

import re
import io
import os
import zipfile
import shutil
import subprocess
from pathlib import Path
from collections import Counter
from typing import Tuple, List, Optional, Dict, Any

# ======= CONFIG =======
CV_DIRECTORY = Path(
    r"C:\Users\BenjaminBusby\OneDrive - MONOCLE SOLUTIONS (PTY) LTD\Documents\Monocle AI Challenge\CV Database"
)
USE_OCR_FALLBACK = False  # set True only if you’ve installed Tesseract + pdf2image

# Precedence thresholds
MIN_LEN_CHARS = 200      # 1) length gate: below this => immediate "too little text"
MAX_LEN_CHARS = 15000
MIN_LANG_CHARS = 120     # min chars required before attempting language detection

# ======= OPTIONAL DEPENDENCIES =======
try:
    import langid
except Exception:
    langid = None

try:
    from pdfminer.high_level import extract_text as pdfminer_extract
except Exception:
    pdfminer_extract = None

try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None

try:
    import docx  # python-docx
except Exception:
    docx = None

try:
    import docx2txt
except Exception:
    docx2txt = None

try:
    from lxml import etree
except Exception:
    etree = None

try:
    import pytesseract
    from pdf2image import convert_from_path
except Exception:
    pytesseract = None
    convert_from_path = None

# ======= Heuristics =======
SECTION_KEYWORDS = {
    "experience", "work experience", "employment", "professional experience",
    "education", "qualifications", "skills", "projects", "certifications",
    "summary", "profile", "objective", "awards", "publications", "university"
}
CV_CLUES = {"curriculum vitae", "resume", "résumé", "cv"}
EMAIL_PAT  = re.compile(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b")
PHONE_PAT  = re.compile(r"\b(?:\+?\d[\d\s().-]{7,}\d)\b")
DATE_PAT   = re.compile(
    r"\b(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)[a-z]*\s+\d{4}\b"
    r"|\b\d{4}\s*[–-]\s*(?:present|\d{4})\b",
    re.IGNORECASE
)
BULLET_PAT = re.compile(r"^\s*[-•\u2022\u25CF\u25E6]", re.MULTILINE)

EN_STOPWORDS = {
    "the","and","to","of","in","for","with","on","as","at","by","from",
    "is","are","was","were","be","been","being",
    "a","an","that","this","these","those","it","its","or","but","if","so",
    "we","our","you","your","they","their","he","she","his","her","them"
}

# ======= Extraction =======
def read_txt(path: Path) -> str:
    for enc in ("utf-8", "cp1252", "latin-1"):
        try:
            return path.read_text(encoding=enc, errors="ignore")
        except Exception:
            continue
    return ""

def read_pdf(path: Path) -> Tuple[str, Optional[int], List[str]]:
    reasons: List[str] = []
    text = ""
    pages = None
    if pdfminer_extract:
        try:
            text = pdfminer_extract(str(path)) or ""
        except Exception as e:
            reasons.append(f"pdfminer error: {e}")
    if not text and PdfReader:
        try:
            r = PdfReader(str(path))
            pages = len(r.pages)
            text = "\n".join((p.extract_text() or "") for p in r.pages)
        except Exception as e:
            reasons.append(f"pypdf error: {e}")
    if not text and USE_OCR_FALLBACK and pytesseract and convert_from_path:
        try:
            images = convert_from_path(str(path), dpi=200)
            text = "\n".join(pytesseract.image_to_string(img) for img in images[:20])
            if not text.strip():
                reasons.append("ocr produced no text")
        except Exception as e:
            reasons.append(f"ocr error: {e}")
    if not text and not reasons:
        reasons.append("no extractable text (maybe scanned/encrypted)")
    return text, pages, reasons

def _docx_xml_texts_from(zipf: zipfile.ZipFile) -> str:
    if not etree:
        return ""
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    texts: List[str] = []
    for name in [n for n in zipf.namelist() if n.startswith("word/") and n.endswith(".xml")]:
        try:
            root = etree.fromstring(zipf.read(name))
        except Exception:
            continue
        texts += root.xpath("//w:t/text()", namespaces=ns)
        texts += root.xpath("//w:txbxContent//w:t/text()", namespaces=ns)
    return "\n".join(t for t in texts if t and t.strip())

def is_real_docx(path: Path) -> bool:
    try:
        with zipfile.ZipFile(str(path), "r"):
            return True
    except Exception:
        return False

# ---- Converters for legacy .doc masquerading as .docx ----
def convert_doc_to_docx_via_word(src: Path, dst: Path) -> bool:
    try:
        import win32com.client  # type: ignore
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(src))
        doc.SaveAs(str(dst), FileFormat=16)  # wdFormatXMLDocument
        doc.Close(False)
        word.Quit()
        return dst.exists() and dst.stat().st_size > 0
    except Exception:
        return False

def find_soffice() -> Optional[str]:
    candidates = [
        "soffice",
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]
    for c in candidates:
        which = shutil.which(c)
        if which:
            return which
        if os.path.exists(c):
            return c
    return None

def convert_doc_to_docx_via_soffice(src: Path, outdir: Path) -> Optional[Path]:
    soffice = find_soffice()
    if not soffice:
        return None
    try:
        outdir.mkdir(parents=True, exist_ok=True)
        subprocess.run(
            [soffice, "--headless", "--convert-to", "docx", str(src), "--outdir", str(outdir)],
            check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE,
        )
        dst = outdir / (src.stem + ".docx")
        return dst if dst.exists() and dst.stat().st_size > 0 else None
    except Exception:
        return None

def ensure_true_docx(path: Path) -> Tuple[Path, Optional[str]]:
    if is_real_docx(path):
        return path, None
    temp_out = path.parent / "_converted_docx"
    temp_out.mkdir(exist_ok=True)
    target = temp_out / (path.stem + ".docx")
    if convert_doc_to_docx_via_word(path, target):
        return target, None
    converted = convert_doc_to_docx_via_soffice(path, temp_out)
    if converted:
        return converted, None
    return path, "conversion failed (no Word/LibreOffice or error during conversion)"

def read_docx(path: Path) -> Tuple[str, List[str]]:
    reasons: List[str] = []
    actual_path = path
    if not is_real_docx(path):
        converted_path, conv_reason = ensure_true_docx(path)
        if conv_reason:
            return "", [f"not a real .docx (bad zip) — {conv_reason}"]
        actual_path = converted_path
    if docx2txt:
        try:
            text = docx2txt.process(str(actual_path)) or ""
            if len(text.strip()) >= 10:
                return text, []
        except Exception as e:
            reasons.append(f"docx2txt error: {e}")
    else:
        reasons.append("docx2txt not installed")
    if docx:
        try:
            d = docx.Document(str(actual_path))
            parts: List[str] = []
            parts.extend(p.text for p in d.paragraphs if p.text and p.text.strip())
            for tbl in d.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        if cell.text and cell.text.strip():
                            parts.append(cell.text)
            for sec in d.sections:
                hdr = getattr(sec, "header", None)
                ftr = getattr(sec, "footer", None)
                if hdr:
                    parts.extend(p.text for p in hdr.paragraphs if p.text and p.text.strip())
                if ftr:
                    parts.extend(p.text for p in ftr.paragraphs if p.text and p.text.strip())
            text = "\n".join(parts)
            if len(text.strip()) >= 10:
                return text, []
            reasons.append("python-docx extracted too little text")
        except Exception as e:
            reasons.append(f"python-docx error: {e}")
    else:
        reasons.append("python-docx not installed")
    try:
        with zipfile.ZipFile(str(actual_path), "r") as z:
            text = _docx_xml_texts_from(z)
            if len(text.strip()) >= 10:
                return text, []
            reasons.append("raw-xml extracted too little text")
    except Exception as e:
        reasons.append(f"raw-xml error: {e}")
    if not reasons:
        reasons.append("docx empty or unreadable")
    return "", reasons

# ======= Language (strict) =======
_hira = re.compile(r"[\u3040-\u309F]")
_kata = re.compile(r"[\u30A0-\u30FF]")
_kanji = re.compile(r"[\u4E00-\u9FFF]")
_jp_punct = re.compile(r"[。、・「」『』（）」]")
_latin = re.compile(r"[A-Za-z]")
_tokenizer = re.compile(r"[A-Za-z']+")

def _script_ratios(s: str) -> dict:
    total_letters = 0
    counts = {"latin": 0, "cjk": 0, "jp_punct": 0}
    for ch in s:
        if _latin.match(ch):
            counts["latin"] += 1; total_letters += 1
        elif _hira.match(ch) or _kata.match(ch) or _kanji.match(ch):
            counts["cjk"] += 1; total_letters += 1
        elif _jp_punct.match(ch):
            counts["jp_punct"] += 1
    total = max(total_letters, 1)
    return {"latin_ratio": counts["latin"]/total, "cjk_ratio": counts["cjk"]/total, "jp_punct": counts["jp_punct"]}

def detect_english_strict(text: str) -> Tuple[bool, str]:
    """Only call if len(text) >= MIN_LANG_CHARS (precedence handled in main)."""
    toks = _tokenizer.findall(text.lower())[:400]
    sw_hits = sum(1 for t in toks if t in EN_STOPWORDS)
    sw_uniq = len({t for t in toks if t in EN_STOPWORDS})
    ratios = _script_ratios(text)
    latin_r, cjk_r, jp_p = ratios["latin_ratio"], ratios["cjk_ratio"], ratios["jp_punct"]
    lid_vote, lid_prob = None, 0.0
    if langid:
        lang, prob = langid.classify(text)
        lid_vote, lid_prob = lang, prob
    if cjk_r >= 0.10 and latin_r <= 0.75:
        return False, f"language: non-English (CJK {cjk_r:.0%}, Latin {latin_r:.0%})"
    if sw_hits < 6 and sw_uniq < 5:
        if not (lid_vote == "en" and lid_prob >= 0.97 and cjk_r < 0.03):
            return False, f"language: non-English (stopwords hits={sw_hits}, unique={sw_uniq})"
    if jp_p >= 5 and sw_hits < 10:
        if not (lid_vote == "en" and lid_prob >= 0.98):
            return False, f"language: non-English (JP punctuation={jp_p}, stopwords={sw_hits})"
    if lid_vote and lid_vote != "en" and lid_prob >= 0.90 and latin_r < 0.85:
        return False, f"language: non-English by model ({lid_vote}, p≈{lid_prob:.2f})"
    return True, ""

# ======= CV checks =======
def looks_like_cv(text: str, threshold: int = 4) -> Tuple[bool, List[str]]:
    t = text.lower()
    signals = 0
    failed: List[str] = []
    if len(text) >= MIN_LEN_CHARS: signals += 1
    else: failed.append("too little text")
    if any(kw in t for kw in SECTION_KEYWORDS): signals += 1
    else: failed.append("no CV section keywords")
    if any(kw in t for kw in CV_CLUES): signals += 1
    else: failed.append("no explicit CV terms")
    if EMAIL_PAT.search(text): signals += 1
    else: failed.append("no email")
    if PHONE_PAT.search(text): signals += 1
    else: failed.append("no phone")
    if DATE_PAT.search(t): signals += 1
    else: failed.append("no employment dates")
    if BULLET_PAT.search(text): signals += 1
    else: failed.append("no bullet points")
    return (signals >= threshold), failed

# ======= Core runner =======
def gatekeep(callback=None) -> Dict[str, Any]:
    """
    Runs the gatekeeper checks, optionally streaming messages via callback(str).
    Returns dict with processed, rejected, error_types, bad_files (list of filenames), bad_stems (list).
    """
    supported = {".pdf", ".docx", ".txt"}
    files = [p for p in CV_DIRECTORY.rglob("*") if p.is_file() and p.suffix.lower() in supported]

    processed = 0
    rejected = 0
    error_types = Counter()
    bad_files: List[str] = []   # filenames (e.g., "CV_0457.pdf")
    bad_stems: List[str] = []   # stems (e.g., "CV_0457")

    def emit(line: str):
        if callback:
            callback(line)
        else:
            print(line)

    for p in sorted(files):
        ext = p.suffix.lower()
        try:
            if ext == ".pdf":
                text, _, pdf_reasons = read_pdf(p)
                extract_fail_reason = "; ".join(pdf_reasons) if not text.strip() else ""
            elif ext == ".docx":
                text, docx_reasons = read_docx(p)
                extract_fail_reason = "; ".join(docx_reasons) if not text.strip() else ""
            else:
                text = read_txt(p)
                extract_fail_reason = "txt: empty or unreadable" if not text.strip() else ""

            # Extraction failure
            if not text.strip():
                emit(f"[{p.name}] No CV detected. Failed tests: extraction: {extract_fail_reason or 'no text'}")
                rejected += 1
                error_types["empty_text"] += 1
                bad_files.append(p.name); bad_stems.append(p.stem)
                continue

            # 1) LENGTH gate
            if len(text) < MIN_LEN_CHARS:
                emit(f"[{p.name}] \nNo CV detected, omitted. Failed tests: too little text (<{MIN_LEN_CHARS} chars)")
                rejected += 1
                error_types["too_short"] += 1
                bad_files.append(p.name); bad_stems.append(p.stem)
                continue
            elif len(text) > MAX_LEN_CHARS:
                emit(f"[{p.name}] \nNo CV detected, omitted. Failed tests: too much text (>{MAX_LEN_CHARS} chars)")
                rejected += 1
                error_types["too_long"] += 1
                bad_files.append(p.name); bad_stems.append(p.stem)
                continue


            # 2) LANGUAGE gate
            if len(text) >= MIN_LANG_CHARS:
                is_en, lang_reason = detect_english_strict(text)
                if not is_en:
                    emit(f"[{p.name}] \nNot in English, omitted. ({lang_reason})")
                    rejected += 1
                    error_types["non_english"] += 1
                    bad_files.append(p.name); bad_stems.append(p.stem)
                    continue

            # 3) CV heuristic
            is_cv, failed = looks_like_cv(text, threshold=4)
            if not is_cv:
                emit(f"[{p.name}] \nNo CV detected. Failed tests: {', '.join(failed)}")
                rejected += 1
                error_types["not_cv"] += 1
                bad_files.append(p.name); bad_stems.append(p.stem)
                continue

            processed += 1

        except Exception as e:
            emit(f"[{p.name}] ERROR: {e}")
            rejected += 1
            error_types["exception"] += 1
            bad_files.append(p.name); bad_stems.append(p.stem)

    emit("\n--- Summary ---")
    emit(f"Processed: {processed}")
    emit(f"Rejected: {rejected}")
    emit(f"Error codes: {dict(error_types)}")

    return {
        "processed": processed,
        "rejected": rejected,
        "error_types": dict(error_types),
        "bad_files": bad_files,
        "bad_stems": bad_stems,
    }

# ======= CLI main =======
def main():
    gatekeep(callback=None)

# ======= GUI hook: redirect print() to callback and return bad_files =======
import contextlib

def run_from_gui(callback):
    """
    Runs gatekeep() and streams every print() line to `callback(str)`.
    Returns the result dict (incl. bad_files, bad_stems) for programmatic use.
    Safe to call from a background thread in GUI.
    """
    class _CallbackWriter:
        def write(self, s):
            s = s.rstrip("\n")
            if s.strip():
                callback(s)
        def flush(self):
            pass

    result_holder = {}
    def _run():
        nonlocal result_holder
        result_holder = gatekeep(callback=callback)

    with contextlib.redirect_stdout(_CallbackWriter()):
        _run()
    return result_holder

def run_and_get_bad_files(callback=None) -> List[str]:
    """
    Convenience: runs gatekeep and returns list of filenames that failed checks.
    If callback is provided, messages are streamed to it.
    """
    res = gatekeep(callback=callback)
    return res.get("bad_files", [])

if __name__ == "__main__":
    main()
